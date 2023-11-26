import json
import os
import xml.etree.ElementTree as ET
import docx
from docx2pdf import convert



def Correcting_report(input_file):
    translate = ['Кодовое имя ОС', 'Краткое имя', 'Имя', 'Домашняя ссылка', 'Цветовой код имени ОС', 'Идентификатор схожих ОС', 'Поддержка', 'Логотип', 'Версия ОС',
                 'Уникальный идентификатор', 'Производитель', 'Имя продукта', 'Версия', 'Серийный номер', 'Идентификатор ОС', 'Тип включения', 'Идентификатор товара', 'Семья']

    original = ['VERSION_CODENAME','PRETTY_NAME', 'NAME', 'HOME_URL', 'ANSI_COLOR', 'ID_LIKE', 'SUPPORT_URL', 'LOGO', 'VERSION_ID',
                'UUID', 'Manufacturer', 'Product Name', 'Version', 'Serial Number', 'ID', 'Wake-up Type', 'SKU Number', 'Family']

    no_need = ['# dmidecode 3.3', 'Getting SMBIOS data from sysfs.', 'SMBIOS 2.5 present.','Handle 0x0001, DMI type 1, 27 bytes', 'System Information']

    # Открываем файл для чтения и чтения данных
    with open(input_file, 'r', encoding='utf-8') as file:
        # Читаем данные из файла
        file_data = file.read()

    for i in range(len(translate)):
        file_data = file_data.replace(original[i], translate[i])

    # Удаляем строки, которые находятся в списке no_need
    lines = file_data.split("\n")
    filtered_lines = [line for line in lines if line.strip() not in no_need]
    new_file_data = "\n".join(filtered_lines)

    with open(f'correct-{input_file}', 'w', encoding='utf-8') as file:
        file.write(new_file_data)

    redact_file = 'correct.txt'



def word_report(redact_file, comp):

    name_key = {'hostname': 'Имя компьютера:', 'uname': 'Версия ядра ОС:', 'cat': 'Информация об операционной системе:',
                'lshw': 'Информация об аппартном обеспечении:',
                'dmidecode': 'Информация о системе и сведения о компьютере:', 'lsusb': 'Подключенные USB-устройства:',
                'apt-mark': 'Список установленного ПО:'}

    with open(redact_file, 'r', encoding='utf-8') as read_file:
        data = json.load(read_file)

    document = docx.Document()
    document.add_heading(f'Отчет {comp}', 0)

    for key, value in data.items():
        if key in ['uname', 'hostname']:
            document.add_heading(f'{name_key[key]}')
            document.add_paragraph(f'{value}')
        else:
            document.add_heading(f'{name_key[key]}')
            document.add_paragraph(f'\n{value}')

    document.save(f'report-{comp}.docx')




def pdf_report(redact_file, comp):
    word_report(redact_file, comp)
    convert(f'report-{comp}.docx')
    os.remove(f'report-{comp}.docx')

def html_report(redact_file, comp):
    name_key = {'hostname': 'Имя компьютера:', 'uname': 'Версия ядра ОС:', 'cat': 'Информация об операционной системе:',
                'lshw': 'Информация об аппартном обеспечении:',
                'dmidecode': 'Информация о системе и сведения о компьютере:', 'lsusb': 'Подключенные USB-устройства:',
                'apt-mark': 'Список установленного ПО:'}

    with open(redact_file, 'r', encoding='utf-8') as file:
        json_data = json.load(file)

    html = "<html>\n<head>\n<style>\nbody {\nbackground-color: #f2f8ff;\n}\n</style>\n</head>\n<body>\n"+f"<h1>Отчет {comp}</h1>\n"

    for key, value in json_data.items():
        html += f"<h2>{name_key[key]}</h2>\n"

        if isinstance(value, dict):
            html += "<ul>\n"
            for k, v in value.items():
                html += f"<li><strong>{k}:</strong> {v}</li>\n"
            html += "</ul>\n"
        else:
            html += f"<p>{value}</p>\n"

    html += "</body>\n</html>"

    with open(f'report-{comp}.html', 'w') as file:
        file.write(html)

def xml_report(redact_file, comp):
    name_key = {'hostname': 'Имя компьютера:', 'uname': 'Версия ядра ОС:', 'cat': 'Информация об операционной системе:',
                'lshw': 'Информация об аппартном обеспечении:',
                'dmidecode': 'Информация о системе и сведения о компьютере:', 'lsusb': 'Подключенные USB-устройства:',
                'apt-mark': 'Список установленного ПО:'}
    # Открываем файл в формате JSON для чтения
    with open(redact_file, 'r', encoding='utf-8') as file:
        # Читаем содержимое файла
        json_data = file.read()

    # Преобразуем данные из формата JSON в словарь Python
    dictionary = json.loads(json_data)

    # Создаем корневой элемент XML и добавляем заголовок отчета
    root = ET.Element('report')
    title = ET.SubElement(root, 'title')
    title.text = f'Отчет {comp}'

    # Проходим по каждому ключу-значению в словаре и создаем соответствующие элементы XML
    for key, value in dictionary.items():
        if key in name_key:
            element = ET.SubElement(root, key)
            element.text = '\n' + name_key[key] + '\n' + str(value) + '\n\n'
        else:
            element = ET.SubElement(root, key)
            element.text = '\n' + str(value) + '\n\n'

    # Создаем объект ElementTree
    tree = ET.ElementTree(root)

    # Сохраняем данные в файл в формате XML
    tree.write(f'report-{comp}.xml', encoding='utf-8', xml_declaration=True)


